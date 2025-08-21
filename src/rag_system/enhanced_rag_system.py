import os
from typing import List, Optional, Dict, Any
import streamlit as st
from llama_index.core import VectorStoreIndex, Document, Settings
import traceback

# Elasticsearch 支援
try:
    from elasticsearch import Elasticsearch
    from llama_index.core.storage.storage_context import StorageContext
    from ..storage.custom_elasticsearch_store import CustomElasticsearchStore
    ELASTICSEARCH_AVAILABLE = True
except ImportError:
    ELASTICSEARCH_AVAILABLE = False

from .rag_system import RAGSystem
from ..storage.conversation_memory import ConversationMemory
from ..processors.user_file_manager import UserFileManager
from ..processors.gemini_ocr import GeminiOCRProcessor
from ..utils.embedding_fix import setup_safe_embedding, prevent_openai_fallback
# from chroma_vector_store import ChromaVectorStoreManager  # 已改用 Elasticsearch

class EnhancedRAGSystem(RAGSystem):
    def __init__(self, use_elasticsearch: bool = True, use_chroma: bool = False):
        super().__init__()
        
        # 初始化新功能模組
        self.memory = ConversationMemory()
        self.file_manager = UserFileManager()
        self.ocr_processor = GeminiOCRProcessor()
        
        # 優先使用 Elasticsearch，停用 ChromaDB
        self.use_elasticsearch = use_elasticsearch
        self.use_chroma = False  # 強制停用 ChromaDB
        self.chroma_manager = None  # 不再使用 ChromaDB
        
        # Elasticsearch 設定
        self.elasticsearch_client = None
        self.elasticsearch_store = None
        
        # 如果啟用 Elasticsearch，嘗試初始化
        if self.use_elasticsearch and ELASTICSEARCH_AVAILABLE:
            self._initialize_elasticsearch()
    
    def _initialize_elasticsearch(self):
        """初始化 Elasticsearch 連接"""
        try:
            from config.config import (
                ELASTICSEARCH_HOST, ELASTICSEARCH_PORT, ELASTICSEARCH_SCHEME,
                ELASTICSEARCH_INDEX_NAME, ELASTICSEARCH_USERNAME, ELASTICSEARCH_PASSWORD,
                ELASTICSEARCH_TIMEOUT, ELASTICSEARCH_MAX_RETRIES, ELASTICSEARCH_VERIFY_CERTS,
                ELASTICSEARCH_VECTOR_DIMENSION
            )
            
            # 建立 Elasticsearch 客戶端
            es_config = {
                'hosts': [f'{ELASTICSEARCH_SCHEME}://{ELASTICSEARCH_HOST}:{ELASTICSEARCH_PORT}'],
                'verify_certs': ELASTICSEARCH_VERIFY_CERTS,
                'ssl_show_warn': False,
                'timeout': ELASTICSEARCH_TIMEOUT,
                'max_retries': ELASTICSEARCH_MAX_RETRIES,
                'retry_on_timeout': True
            }
            
            if ELASTICSEARCH_USERNAME and ELASTICSEARCH_PASSWORD:
                es_config['basic_auth'] = (ELASTICSEARCH_USERNAME, ELASTICSEARCH_PASSWORD)
            
            self.elasticsearch_client = Elasticsearch(**es_config)
            print(f"🔧 EnhancedRAGSystem ES客戶端類型: {type(self.elasticsearch_client)}")
            
            # 檢查連接
            if self.elasticsearch_client.ping():
                st.info("✅ Elasticsearch 連接成功")
                # 維度一致性檢查
                if not self._validate_embedding_dimension(ELASTICSEARCH_VECTOR_DIMENSION):
                    st.error("❌ 嵌入維度與 Elasticsearch 配置不一致，請檢查設定。")
                    self.use_elasticsearch = False
                    return False

                # 建立 vector store（使用自定義實現避免 async/await 問題）
                self.elasticsearch_store = CustomElasticsearchStore(
                    es_client=self.elasticsearch_client,
                    index_name=ELASTICSEARCH_INDEX_NAME,
                    vector_field="embedding",
                    text_field="content",
                    metadata_field="metadata"
                )
                return True
            else:
                st.error("❌ 無法連接到 Elasticsearch，系統無法正常運行")
                self.use_elasticsearch = False
                return False
                
        except Exception as e:
            st.error(f"❌ Elasticsearch 初始化失敗: {str(e)}，系統無法正常運行")
            self.use_elasticsearch = False
            return False
    
    def _ensure_models_initialized(self):
        """確保模型已初始化"""
        if not self.models_initialized:
            self._setup_models()
            self.models_initialized = True
    
    def _setup_models(self):
        """設定模型 - 覆寫父類方法以確保正確初始化"""
        from config.config import GROQ_API_KEY, LLM_MODEL, JINA_API_KEY
        from llama_index.llms.groq import Groq
        from llama_index.core.node_parser import SimpleNodeParser
        from llama_index.core import Settings
        import streamlit as st
        
        # 防止 OpenAI 回退並設置安全的嵌入模型
        prevent_openai_fallback()
        
        # 設定LLM
        if GROQ_API_KEY:
            llm = Groq(model=LLM_MODEL, api_key=GROQ_API_KEY)
        else:
            st.error("請設定GROQ_API_KEY環境變數")
            return
        
        # 設定安全嵌入模型 - 僅使用 Jina API（含本地後備）
        try:
            embed_model = setup_safe_embedding(JINA_API_KEY)
            st.success("✅ 成功初始化嵌入模型（Jina）")
        except Exception as e2:
            st.error(f"嵌入模型初始化失敗: {str(e2)}")
            return
        
        # 設定全域配置
        Settings.llm = llm
        Settings.embed_model = embed_model
        Settings.node_parser = SimpleNodeParser.from_defaults(chunk_size=1024)
        
        st.success("🔧 模型初始化完成")
        
    def _get_embed_dim(self) -> int:
        """嘗試從當前嵌入模型取得維度。找不到則返回 None。"""
        try:
            model = Settings.embed_model
            for attr in ("embed_dim", "_embed_dim", "dimension", "dim"):
                if hasattr(model, attr):
                    val = getattr(model, attr)
                    try:
                        return int(val)
                    except Exception:
                        continue
        except Exception:
            pass
        return None

    def _validate_embedding_dimension(self, expected_dim: int) -> bool:
        """驗證當前嵌入模型維度與預期一致。"""
        actual = self._get_embed_dim()
        if actual is None:
            st.warning("無法檢測嵌入維度，跳過維度驗證。")
            return True
        if actual != int(expected_dim):
            st.error(f"嵌入維度不匹配：模型為 {actual}，Elasticsearch 預期為 {expected_dim}")
            return False
        return True

    def query_with_context(self, question: str) -> str:
        """帶上下文記憶的查詢"""
        if not self.query_engine:
            return "系統尚未初始化，請先載入文件。"
        
        try:
            # 建構包含歷史對話的完整查詢
            context_prompt = self.memory.get_context_prompt()
            
            if context_prompt and self.memory.is_enabled():
                enhanced_question = f"""
{context_prompt}

當前問題: {question}

請基於以上對話歷史和知識庫內容回答當前問題。如果當前問題與之前的對話相關，請考慮上下文語境。
"""
            else:
                enhanced_question = question
            
            with st.spinner("正在思考您的問題..."):
                response = self.query_engine.query(enhanced_question)
                response_str = str(response)
                
                # 將這輪對話加入記憶
                self.memory.add_exchange(question, response_str)
                
                return response_str
                
        except Exception as e:
            st.error(f"查詢時發生錯誤: {str(e)}")
            st.write(traceback.format_exc())
            return "抱歉，處理您的問題時發生錯誤。"
    
    def process_uploaded_files(self, uploaded_files) -> List[Document]:
        """處理上傳的檔案"""
        import logging
        logging.basicConfig(level=logging.INFO)
        logger = logging.getLogger(__name__)
        
        logger.info(f"🚀 開始處理上傳文件，共 {len(uploaded_files) if uploaded_files else 0} 個文件")
        
        if not uploaded_files:
            logger.warning("⚠️ 沒有上傳的文件")
            return []
        
        documents = []
        
        for i, uploaded_file in enumerate(uploaded_files):
            logger.info(f"📄 處理文件 {i+1}/{len(uploaded_files)}: {uploaded_file.name}")
            logger.info(f"   - 文件大小: {uploaded_file.size:,} bytes ({uploaded_file.size/(1024*1024):.2f} MB)")
            logger.info(f"   - 文件類型: {uploaded_file.type if hasattr(uploaded_file, 'type') else '未知'}")
            
            try:
                # 儲存檔案
                logger.info(f"💾 正在儲存文件: {uploaded_file.name}")
                file_path = self.file_manager.save_uploaded_file(uploaded_file)
                
                if not file_path:
                    logger.error(f"❌ 文件儲存失敗: {uploaded_file.name}")
                    continue
                    
                logger.info(f"✅ 文件儲存成功: {file_path}")
                
                # 根據檔案類型處理
                if self.file_manager.is_image_file(uploaded_file.name):
                    logger.info(f"🖼️ 處理圖片文件: {uploaded_file.name}")
                    # 圖片OCR處理
                    doc = self._process_image_file(uploaded_file, file_path)
                elif self.file_manager.is_document_file(uploaded_file.name):
                    logger.info(f"📝 處理文檔文件: {uploaded_file.name}")
                    # 文檔處理
                    doc = self._process_document_file(uploaded_file, file_path)
                else:
                    logger.warning(f"❓ 不支援的檔案類型: {uploaded_file.name}")
                    st.warning(f"不支援的檔案類型: {uploaded_file.name}")
                    continue
                
                if doc:
                    logger.info(f"✅ 文件處理成功，生成文檔: {uploaded_file.name}")
                    documents.append(doc)
                else:
                    logger.error(f"❌ 文件處理失敗，沒有生成文檔: {uploaded_file.name}")
                    
            except Exception as e:
                logger.error(f"❌ 處理檔案 {uploaded_file.name} 時發生錯誤: {str(e)}")
                logger.error(f"   詳細錯誤信息: {type(e).__name__}: {str(e)}")
                import traceback
                logger.error(f"   錯誤堆疊: {traceback.format_exc()}")
                st.error(f"處理檔案 {uploaded_file.name} 時發生錯誤: {str(e)}")
                continue
        
        logger.info(f"🎉 文件處理完成，成功處理 {len(documents)}/{len(uploaded_files)} 個文件")
        return documents
    
    def _process_image_file(self, uploaded_file, file_path: str) -> Optional[Document]:
        """處理圖片檔案"""
        import logging
        logger = logging.getLogger(__name__)
        
        logger.info(f"🖼️ 開始處理圖片文件: {uploaded_file.name}")
        
        if not self.ocr_processor.is_available():
            logger.warning(f"⚠️ OCR服務不可用，跳過圖片檔案: {uploaded_file.name}")
            st.warning(f"OCR服務不可用，跳過圖片檔案: {uploaded_file.name}")
            return None
        
        try:
            logger.info(f"📖 正在讀取圖片數據: {uploaded_file.name}")
            # 讀取圖片數據
            image_data = self.file_manager.get_file_content(os.path.basename(file_path))
            if not image_data:
                logger.error(f"❌ 無法讀取圖片數據: {uploaded_file.name}")
                return None
            
            logger.info(f"   - 圖片數據大小: {len(image_data):,} bytes")
            
            # 取得檔案擴展名
            file_ext = os.path.splitext(uploaded_file.name)[1].lower().lstrip('.')
            logger.info(f"   - 檔案格式: {file_ext}")
            
            # OCR處理
            logger.info(f"🔍 開始OCR處理: {uploaded_file.name}")
            with st.spinner(f"正在進行OCR處理: {uploaded_file.name}"):
                ocr_result = self.ocr_processor.extract_text_from_image(image_data, file_ext)
            
            logger.info(f"   - OCR處理完成，成功: {ocr_result['success']}")
            
            if ocr_result['success']:
                text_length = len(ocr_result['text'])
                logger.info(f"   - 提取的文本長度: {text_length} 字符")
                logger.info(f"   - OCR信心度: {ocr_result.get('confidence', 'unknown')}")
                
                # 建立文檔
                document = Document(
                    text=ocr_result['text'],
                    metadata={
                        "source": uploaded_file.name,
                        "type": "image_ocr",
                        "original_format": file_ext,
                        "file_size": uploaded_file.size,
                        "ocr_confidence": ocr_result.get('confidence', 'unknown'),
                        "processed_at": st.session_state.get('current_time', 'unknown')
                    }
                )
                
                logger.info(f"✅ OCR處理成功: {uploaded_file.name}")
                st.success(f"✅ OCR處理成功: {uploaded_file.name}")
                return document
            else:
                error_msg = ocr_result.get('error', '未知錯誤')
                logger.error(f"❌ OCR處理失敗: {uploaded_file.name} - {error_msg}")
                st.error(f"❌ OCR處理失敗: {uploaded_file.name} - {error_msg}")
                return None
                
        except Exception as e:
            logger.error(f"❌ 處理圖片檔案時發生錯誤: {uploaded_file.name} - {str(e)}")
            import traceback
            logger.error(f"   錯誤堆疊: {traceback.format_exc()}")
            st.error(f"處理圖片檔案時發生錯誤: {str(e)}")
            return None
    
    def _process_document_file(self, uploaded_file, file_path: str) -> Optional[Document]:
        """處理文檔檔案"""
        import logging
        logger = logging.getLogger(__name__)
        
        logger.info(f"📝 開始處理文檔文件: {uploaded_file.name}")
        
        try:
            file_ext = os.path.splitext(uploaded_file.name)[1].lower()
            logger.info(f"   - 檔案格式: {file_ext}")
            logger.info(f"   - 檔案路徑: {file_path}")
            
            if file_ext == '.pdf':
                logger.info(f"📄 開始PDF處理: {uploaded_file.name}")
                # PDF處理
                docs = self.load_pdfs([file_path])
                if docs:
                    doc = docs[0]
                    logger.info(f"   - PDF解析成功，文本長度: {len(doc.text)} 字符")
                    # 更新元數據
                    doc.metadata.update({
                        "type": "user_document",
                        "file_size": uploaded_file.size,
                        "uploaded_at": st.session_state.get('current_time', 'unknown')
                    })
                    logger.info(f"✅ PDF處理完成: {uploaded_file.name}")
                    return doc
                else:
                    logger.error(f"❌ PDF處理失敗，無法解析: {uploaded_file.name}")
                    return None
            
            elif file_ext == '.txt':
                logger.info(f"📄 開始TXT處理: {uploaded_file.name}")
                # 文字檔處理
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                logger.info(f"   - TXT解析成功，文本長度: {len(text)} 字符")
                document = Document(
                    text=text,
                    metadata={
                        "source": uploaded_file.name,
                        "type": "user_document",
                        "original_format": "txt",
                        "file_size": uploaded_file.size,
                        "uploaded_at": st.session_state.get('current_time', 'unknown')
                    }
                )
                logger.info(f"✅ TXT處理完成: {uploaded_file.name}")
                return document
            
            elif file_ext in ['.md', '.markdown']:
                logger.info(f"📝 開始Markdown處理: {uploaded_file.name}")
                # Markdown檔處理
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
                
                logger.info(f"   - Markdown解析成功，文本長度: {len(text)} 字符")
                document = Document(
                    text=text,
                    metadata={
                        "source": uploaded_file.name,
                        "type": "user_document",
                        "original_format": "markdown",
                        "file_size": uploaded_file.size,
                        "uploaded_at": st.session_state.get('current_time', 'unknown')
                    }
                )
                logger.info(f"✅ Markdown處理完成: {uploaded_file.name}")
                return document
            
            elif file_ext == '.docx':
                logger.info(f"📄 開始DOCX處理: {uploaded_file.name}")
                # DOCX檔處理
                try:
                    import docx
                    doc = docx.Document(file_path)
                    text = ""
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    logger.info(f"   - DOCX解析成功，文本長度: {len(text)} 字符")
                    document = Document(
                        text=text,
                        metadata={
                            "source": uploaded_file.name,
                            "type": "user_document",
                            "original_format": "docx",
                            "file_size": uploaded_file.size,
                            "uploaded_at": st.session_state.get('current_time', 'unknown')
                        }
                    )
                    logger.info(f"✅ DOCX處理完成: {uploaded_file.name}")
                    return document
                    
                except ImportError as e:
                    logger.error(f"❌ DOCX處理失敗，缺少依賴: {uploaded_file.name} - {str(e)}")
                    st.error("需要安裝 python-docx 套件來處理 DOCX 檔案")
                    return None
                except Exception as e:
                    logger.error(f"❌ DOCX處理失敗: {uploaded_file.name} - {str(e)}")
                    st.error(f"DOCX 檔案處理失敗: {str(e)}")
                    return None
            
            else:
                logger.warning(f"❓ 暫不支援的文檔格式: {file_ext} - {uploaded_file.name}")
                st.warning(f"暫不支援的文檔格式: {file_ext}")
                return None
                
        except Exception as e:
            logger.error(f"❌ 處理文檔檔案時發生錯誤: {uploaded_file.name} - {str(e)}")
            import traceback
            logger.error(f"   錯誤堆疊: {traceback.format_exc()}")
            st.error(f"處理文檔檔案時發生錯誤: {str(e)}")
            return None
    
    def create_index(self, documents: List[Document]) -> VectorStoreIndex:
        """建立新的向量索引 (優先使用 Elasticsearch)"""
        with st.spinner("正在建立向量索引..."):
            # 確保模型已正確初始化
            self._ensure_models_initialized()
            
            index = None
            
            try:
                # 優先使用 Elasticsearch
                if self.use_elasticsearch and self.elasticsearch_store:
                    st.info("使用 Elasticsearch 建立索引...")
                    try:
                        # 建立前做維度驗證
                        from config.config import ELASTICSEARCH_VECTOR_DIMENSION
                        if not self._validate_embedding_dimension(ELASTICSEARCH_VECTOR_DIMENSION):
                            st.error("❌ 維度不一致，停止建立索引。")
                            return None

                        # 創建索引 - 增加詳細日誌
                        print(f"🚀 開始使用 ES 建立索引，文檔數量: {len(documents)}")
                        st.info(f"📊 準備向量化 {len(documents)} 個文檔")
                        
                        # 檢查文檔內容
                        for i, doc in enumerate(documents[:3]):  # 只檢查前3個
                            content_preview = doc.text[:100] + "..." if len(doc.text) > 100 else doc.text
                            print(f"📄 文檔 {i+1}: {len(doc.text)} 字符")
                            print(f"   內容預覽: {content_preview}")
                            if hasattr(doc, 'metadata') and doc.metadata:
                                print(f"   元數據: {doc.metadata}")
                        
                        # 建立 storage context
                        storage_context = StorageContext.from_defaults(
                            vector_store=self.elasticsearch_store
                        )
                        
                        # 先創建空索引，然後逐個添加文檔以避免 async 問題
                        index = VectorStoreIndex([], storage_context=storage_context)
                        
                        # 逐個添加文檔到索引
                        st.info("正在逐個添加文檔到索引...")
                        progress_bar = st.progress(0)
                        for i, doc in enumerate(documents):
                            try:
                                index.insert(doc)
                                progress_bar.progress((i + 1) / len(documents))
                            except Exception as doc_error:
                                st.warning(f"文檔 {i+1} 添加失敗: {str(doc_error)}")
                                continue
                        progress_bar.empty()
                        
                        # 強制刷新 ES 索引
                        if hasattr(self, 'elasticsearch_client') and self.elasticsearch_client:
                            try:
                                print(f"🔄 EnhancedRAGSystem刷新ES索引，客戶端類型: {type(self.elasticsearch_client)}")
                                # 使用正確的索引名稱
                                index_name = getattr(self, 'index_name', None)
                                if not index_name and hasattr(self, 'elasticsearch_store'):
                                    index_name = getattr(self.elasticsearch_store, 'index_name', 'rag_intelligent_assistant')
                                self.elasticsearch_client.indices.refresh(index=index_name)
                                print("✅ ES索引已刷新")
                                
                                # 驗證索引結果
                                stats = self.elasticsearch_client.indices.stats(index=index_name)
                                doc_count = stats['indices'][index_name]['total']['docs']['count']
                                print(f"📊 索引驗證: {doc_count} 個文檔已索引")
                                st.info(f"📊 已成功索引 {doc_count} 個文檔到 Elasticsearch")
                                
                            except Exception as refresh_error:
                                print(f"⚠️ 索引刷新或驗證失敗: {refresh_error}")
                        
                        st.success("✅ 成功使用 Elasticsearch 建立索引")
                        
                    except Exception as e:
                        st.warning(f"Elasticsearch 索引創建失敗: {str(e)}")
                        st.error("❌ Elasticsearch 索引創建失敗，系統無法正常運行")
                        return None
                
                if index:
                    self.index = index
                    return index
                else:
                    st.error("索引創建失敗")
                    return None
                    
            except Exception as e:
                st.error(f"建立索引時發生未預期錯誤: {str(e)}")
                return None
    
    def load_existing_index(self) -> bool:
        """載入現有的向量索引 (優先使用 Elasticsearch)"""
        # 確保模型已初始化
        self._ensure_models_initialized()
        
        try:
            # 優先嘗試 Elasticsearch
            if self.use_elasticsearch and self.elasticsearch_store:
                st.info("嘗試從 Elasticsearch 載入索引...")
                try:
                    # 載入前做維度驗證
                    from config.config import ELASTICSEARCH_VECTOR_DIMENSION
                    if not self._validate_embedding_dimension(ELASTICSEARCH_VECTOR_DIMENSION):
                        st.error("❌ 維度不一致，停止載入索引。")
                        self.use_elasticsearch = False
                        # Elasticsearch 維度不一致，停止載入
                    
                    # 檢查 Elasticsearch 是否有資料
                    es_stats = self.elasticsearch_client.indices.stats(
                        index=self.elasticsearch_store.index_name
                    )
                    doc_count = es_stats['indices'][self.elasticsearch_store.index_name]['total']['docs']['count']
                    
                    if doc_count > 0:
                        # 從 Elasticsearch 重建索引 - 使用同步方式
                        storage_context = StorageContext.from_defaults(
                            vector_store=self.elasticsearch_store
                        )
                        # 直接創建索引實例而不使用 from_vector_store
                        self.index = VectorStoreIndex(
                            nodes=[],
                            storage_context=storage_context
                        )
                        self.setup_query_engine()
                        st.success(f"✅ 成功從 Elasticsearch 載入 {doc_count} 個文檔")
                        return True
                    else:
                        st.info("Elasticsearch 索引為空")
                        
                except Exception as e:
                    st.warning(f"Elasticsearch 載入失敗: {str(e)}")
                    self.use_elasticsearch = False
            
            # 如果 Elasticsearch 載入失敗，系統無法正常運行
            st.error("❌ Elasticsearch 載入失敗，系統無法正常運行")
            return False
                
        except Exception as e:
            st.error(f"載入索引時發生未預期錯誤: {str(e)}")
            return False
    
    
    

    def rebuild_index_with_user_files(self) -> bool:
        """重建索引，包含用戶上傳的檔案"""
        try:
            all_documents = []
            
            # 載入官方茶葉資料
            from ..processors.enhanced_pdf_downloader import EnhancedPDFDownloader
            downloader = EnhancedPDFDownloader()
            official_pdfs = downloader.get_existing_pdfs()
            
            if official_pdfs:
                st.info("載入官方茶葉資料...")
                official_docs = self.load_pdfs(official_pdfs)
                # 標記為官方資料
                for doc in official_docs:
                    doc.metadata["data_source"] = "official"
                all_documents.extend(official_docs)
            
            # 載入用戶上傳的檔案
            user_files = self.file_manager.get_uploaded_files()
            if user_files:
                st.info("載入用戶上傳的檔案...")
                
                # 模擬uploaded_file對象來重用現有邏輯
                class MockUploadedFile:
                    def __init__(self, file_info):
                        self.name = file_info['name']
                        self.size = file_info['size']
                
                mock_files = [MockUploadedFile(f) for f in user_files]
                user_docs = []
                
                for i, file_info in enumerate(user_files):
                    if file_info['type'] == 'image':
                        doc = self._process_image_file(mock_files[i], file_info['path'])
                    else:
                        doc = self._process_document_file(mock_files[i], file_info['path'])
                    
                    if doc:
                        doc.metadata["data_source"] = "user_upload"
                        user_docs.append(doc)
                
                all_documents.extend(user_docs)
            
            if all_documents:
                # 重建索引
                st.info("重建向量索引...")
                index = self.create_index(all_documents)
                
                if index:
                    self.setup_query_engine()
                    return True
            
            return False
            
        except Exception as e:
            st.error(f"重建索引時發生錯誤: {str(e)}")
            return False
    
    def get_document_statistics(self) -> dict:
        """取得文件統計資訊 (僅支援 Elasticsearch)"""
        if not self.index:
            return {}
        
        stats = {
            "total_documents": 0,
            "total_nodes": 0,
            "document_details": [],
            "total_pages": 0
        }
        
        try:
            if self.use_elasticsearch and self.elasticsearch_client:
                # 使用 Elasticsearch 統計
                try:
                    from config.config import ELASTICSEARCH_INDEX_NAME
                    index_name = ELASTICSEARCH_INDEX_NAME or 'rag_intelligent_assistant'
                    
                    es_stats = self.elasticsearch_client.indices.stats(
                        index=index_name
                    )
                    doc_count = es_stats['indices'][index_name]['total']['docs']['count']
                    index_size = es_stats['indices'][index_name]['total']['store']['size_in_bytes']
                    
                    stats["total_documents"] = doc_count
                    stats["total_nodes"] = doc_count
                    stats["index_size_bytes"] = index_size
                    stats["index_size_mb"] = round(index_size / 1024 / 1024, 2)
                    
                    # 從 Elasticsearch 獲取文檔類型統計
                    print(f"🔍 EnhancedRAGSystem執行ES搜尋，客戶端類型: {type(self.elasticsearch_client)}")
                    search_result = self.elasticsearch_client.search(
                        index=index_name,
                        body={
                            "size": 0,
                            "aggs": {
                                "source_types": {
                                    "terms": {
                                        "field": "metadata.source.keyword",
                                        "size": 100
                                    }
                                }
                            }
                        }
                    )
                    print(f"✅ EnhancedRAGSystem ES查詢響應類型: {type(search_result)}")
                    if hasattr(search_result, '__await__'):
                        print("🚨 EnhancedRAGSystem檢測到awaitable response - 異步客戶端錯誤！")
                        raise Exception("EnhancedRAGSystem同步客戶端返回了awaitable response")
                    
                    source_buckets = search_result.get('aggregations', {}).get('source_types', {}).get('buckets', [])
                    for bucket in source_buckets:
                        stats["document_details"].append({
                            "name": bucket['key'],
                            "pages": bucket['doc_count'],
                            "node_count": bucket['doc_count']
                        })
                    
                    st.info(f"📊 從 Elasticsearch 獲取統計: {stats['total_documents']} 個文檔")
                    
                except Exception as es_e:
                    error_msg = str(es_e)
                    print(f"❌ EnhancedRAGSystem ES統計錯誤: {error_msg}")
                    print(f"🔧 錯誤類型: {type(es_e)}")
                    if "ObjectApiResponse" in error_msg or "await" in error_msg or "coroutine" in error_msg:
                        print("🚨 EnhancedRAGSystem檢測到ObjectApiResponse錯誤！")
                        print(f"🔧 當前ES客戶端類型: {type(self.elasticsearch_client)}")
                    import traceback
                    print(f"🔍 EnhancedRAGSystem完整錯誤堆疊: {traceback.format_exc()}")
                    st.error(f"無法從 Elasticsearch 獲取統計資訊: {str(es_e)}")
            else:
                st.error("❌ Elasticsearch 後端未啟用，系統無法獲取統計資訊")
            
        except Exception as e:
            st.error(f"獲取文檔統計時發生錯誤: {str(e)}")
        
        return stats
    
    def get_indexed_files(self) -> List[Dict[str, Any]]:
        """獲取已索引的文件列表（僅支援 Elasticsearch）"""
        files = []
        
        try:
            if self.use_elasticsearch and self.elasticsearch_store:
                # 從 Elasticsearch 獲取文件列表
                files = self._get_elasticsearch_files()
            else:
                st.warning("⚠️ 只支援 Elasticsearch 後端，請確保 Elasticsearch 已正確配置")
        except Exception as e:
            st.error(f"獲取文件列表時發生錯誤: {str(e)}")
        
        return files
    
    def _get_elasticsearch_files(self) -> List[Dict[str, Any]]:
        """從 Elasticsearch 獲取文件列表"""
        files = []
        try:
            # 搜索所有文檔
            response = self.elasticsearch_client.search(
                index=self.elasticsearch_store.index_name,
                body={
                    "query": {"match_all": {}},
                    "size": 1000,
                    "_source": ["metadata", "content"]
                }
            )
            
            file_map = {}
            for hit in response['hits']['hits']:
                metadata = hit['_source'].get('metadata', {})
                source = metadata.get('source', '未知文件')
                
                if source not in file_map:
                    file_map[source] = {
                        'id': source,
                        'name': os.path.basename(source),
                        'path': source,
                        'type': metadata.get('file_type', 'unknown'),
                        'upload_time': metadata.get('upload_time', '未知'),
                        'size': metadata.get('file_size', 0),
                        'page_count': metadata.get('pages', 1),
                        'node_count': 0
                    }
                
                file_map[source]['node_count'] += 1
            
            files = list(file_map.values())
            
        except Exception as e:
            st.error(f"從 Elasticsearch 獲取文件列表失敗: {str(e)}")
        
        return files
    
    def delete_file_from_knowledge_base(self, file_id: str) -> bool:
        """從知識庫中刪除指定文件"""
        try:
            success = False
            
            if self.use_elasticsearch and self.elasticsearch_store:
                # 從 Elasticsearch 刪除文件
                success = self._delete_from_elasticsearch(file_id)
            else:
                st.error("❌ 只支援 Elasticsearch 後端刪除操作")
                success = False
            
            if success:
                # 同時從文件系統刪除（如果存在）
                self._delete_from_filesystem(file_id)
                st.success(f"✅ 文件 {os.path.basename(file_id)} 已從知識庫中刪除")
            
            return success
            
        except Exception as e:
            st.error(f"刪除文件時發生錯誤: {str(e)}")
            return False
    
    def _delete_from_elasticsearch(self, file_id: str) -> bool:
        """從 Elasticsearch 刪除文件的所有節點"""
        try:
            # 查詢該文件的所有文檔
            response = self.elasticsearch_client.search(
                index=self.elasticsearch_store.index_name,
                body={
                    "query": {
                        "term": {
                            "metadata.source.keyword": file_id
                        }
                    },
                    "size": 1000
                }
            )
            
            # 刪除所有相關文檔
            for hit in response['hits']['hits']:
                self.elasticsearch_client.delete(
                    index=self.elasticsearch_store.index_name,
                    id=hit['_id']
                )
            
            # 刷新索引
            self.elasticsearch_client.indices.refresh(
                index=self.elasticsearch_store.index_name
            )
            
            return True
            
        except Exception as e:
            st.error(f"從 Elasticsearch 刪除文件失敗: {str(e)}")
            return False
    
    def _delete_from_filesystem(self, file_id: str):
        """從文件系統刪除文件（如果存在）"""
        try:
            if os.path.exists(file_id):
                os.remove(file_id)
                print(f"✅ 已從文件系統刪除: {file_id}")
        except Exception as e:
            print(f"⚠️ 從文件系統刪除文件失敗: {str(e)}")
    
    def get_enhanced_statistics(self) -> Dict[str, Any]:
        """取得增強的統計資訊"""
        base_stats = self.get_document_statistics()
        
        # 用戶檔案統計
        file_stats = self.file_manager.get_file_stats()
        
        # 對話記憶統計
        memory_stats = self.memory.get_memory_stats()
        
        # OCR可用性
        ocr_available = self.ocr_processor.is_available()
        
        return {
            "base_statistics": base_stats,
            "user_files": file_stats,
            "conversation_memory": memory_stats,
            "ocr_available": ocr_available,
            "total_data_sources": {
                "official_documents": base_stats.get("total_documents", 0) - file_stats.get("total_files", 0),
                "user_documents": file_stats.get("document_count", 0),
                "user_images": file_stats.get("image_count", 0)
            }
        }
    
    def clear_conversation_memory(self):
        """清除對話記憶"""
        self.memory.clear_memory()
    
    def get_memory_for_display(self):
        """取得用於顯示的記憶內容"""
        return self.memory.get_memory_for_display()
    
    def delete_user_file(self, filename: str) -> bool:
        """刪除用戶檔案"""
        return self.file_manager.delete_file(filename)
